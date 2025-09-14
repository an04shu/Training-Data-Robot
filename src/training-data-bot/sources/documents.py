"""
Document loader for text-based formats.

This module handles loading of various text-based document formats
including TXT, Markdown, HTML, JSON, and CSV files.
"""

import json
import csv
from pathlib import Path
from typing import Union, Dict, Any, List
import asyncio

from .base import BaseLoader
from ..core.models import Document, DocumentType
from ..core.exceptions import DocumentLoadError
from ..core.logging import LogContext


class DocumentLoader(BaseLoader):
    """
    Loader for text-based document formats.
    
    Supports: TXT, MD, HTML, JSON, CSV, DOCX
    """

    def __init__(self):
        super().__init__()
        self.supported_formats = [
            DocumentType.TXT,
            DocumentType.MD,
            DocumentType.HTML,
            DocumentType.JSON,
            DocumentType.CSV,
            DocumentType.DOCX,
        ]

    async def load_single(
        self,
        source: Union[str, Path],
        encoding: str = "utf-8",
        **kwargs
    ) -> Document:
        """
        Load a single text-based document.
        
        Args:
            source: File path
            encoding: Text encoding
            **kwargs: Additional options
            
        Returns:
            Loaded document
        """
        source = Path(source)
        
        if not source.exists():
            raise DocumentLoadError(f"File not found: {source}")
        
        doc_type = self.get_document_type(source)
        
        with LogContext("load_document", file=str(source), type=doc_type.value):
            try:
                # Route to appropriate loader method
                if doc_type == DocumentType.TXT:
                    content = await self._load_text(source, encoding)
                elif doc_type == DocumentType.MD:
                    content = await self._load_markdown(source, encoding)
                elif doc_type == DocumentType.HTML:
                    content = await self._load_html(source, encoding)
                elif doc_type == DocumentType.JSON:
                    content = await self._load_json(source, encoding)
                elif doc_type == DocumentType.CSV:
                    content = await self._load_csv(source, encoding)
                elif doc_type == DocumentType.DOCX:
                    content = await self._load_docx(source)
                else:
                    raise DocumentLoadError(f"Unsupported format: {doc_type}")
                
                # Create document
                title = source.stem
                document = self.create_document(
                    title=title,
                    content=content,
                    source=source,
                    doc_type=doc_type,
                    encoding=encoding,
                    extraction_method=f"DocumentLoader.{doc_type.value}",
                )
                
                return document
                
            except Exception as e:
                raise DocumentLoadError(
                    f"Failed to load {doc_type.value} file: {source}",
                    file_path=str(source),
                    cause=e
                )

    async def _load_text(self, path: Path, encoding: str) -> str:
        """Load plain text file."""
        return await asyncio.to_thread(path.read_text, encoding=encoding)

    async def _load_markdown(self, path: Path, encoding: str) -> str:
        """Load Markdown file."""
        # For now, treat as plain text
        # In a full implementation, you might convert to HTML or extract metadata
        return await asyncio.to_thread(path.read_text, encoding=encoding)

    async def _load_html(self, path: Path, encoding: str) -> str:
        """Load HTML file and extract text content."""
        def _extract_html_text():
            try:
                from bs4 import BeautifulSoup
                
                with open(path, 'r', encoding=encoding) as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text
                
            except ImportError:
                # Fallback: read as plain text if BeautifulSoup not available
                return path.read_text(encoding=encoding)
        
        return await asyncio.to_thread(_extract_html_text)

    async def _load_json(self, path: Path, encoding: str) -> str:
        """Load JSON file and convert to readable text."""
        def _convert_json():
            with open(path, 'r', encoding=encoding) as f:
                data = json.load(f)
            
            # Convert JSON to readable text format
            if isinstance(data, dict):
                lines = []
                for key, value in data.items():
                    if isinstance(value, (dict, list)):
                        lines.append(f"{key}: {json.dumps(value, indent=2)}")
                    else:
                        lines.append(f"{key}: {value}")
                return "\n".join(lines)
            elif isinstance(data, list):
                lines = []
                for i, item in enumerate(data):
                    lines.append(f"Item {i+1}: {json.dumps(item, indent=2)}")
                return "\n".join(lines)
            else:
                return str(data)
        
        return await asyncio.to_thread(_convert_json)

    async def _load_csv(self, path: Path, encoding: str) -> str:
        """Load CSV file and convert to readable text."""
        def _convert_csv():
            lines = []
            with open(path, 'r', encoding=encoding, newline='') as f:
                reader = csv.reader(f)
                headers = next(reader, None)
                
                if headers:
                    lines.append("Headers: " + ", ".join(headers))
                    lines.append("")
                
                for row_num, row in enumerate(reader, 1):
                    if headers and len(row) == len(headers):
                        row_data = []
                        for header, value in zip(headers, row):
                            row_data.append(f"{header}: {value}")
                        lines.append(f"Row {row_num}: {' | '.join(row_data)}")
                    else:
                        lines.append(f"Row {row_num}: {', '.join(row)}")
                    
                    # Limit rows to prevent excessive memory usage
                    if row_num > 1000:
                        lines.append(f"... (truncated, total rows exceed 1000)")
                        break
            
            return "\n".join(lines)
        
        return await asyncio.to_thread(_convert_csv)

    async def _load_docx(self, path: Path) -> str:
        """Load DOCX file and extract text."""
        def _extract_docx_text():
            try:
                from docx import Document
                
                doc = Document(path)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                return "\n".join(text_parts)
                
            except ImportError:
                raise DocumentLoadError(
                    "python-docx package required for DOCX files. Install with: pip install python-docx"
                )
        
        return await asyncio.to_thread(_extract_docx_text) 